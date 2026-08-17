#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REPO_ROOT = Path(__file__).resolve().parents[1]
PRODUCTS_PATH = REPO_ROOT / "public/data/products.json"
COLLECTION_PRODUCTS_PATH = REPO_ROOT / "public/data/collection-products.json"
OUTPUT_DOCX = REPO_ROOT / "output/salt-aeo-geo-resource-hub-content.docx"

RESOURCE_BASE = "/resources"
PRODUCT_BASE = "/products"

NAVY = RGBColor(15, 39, 66)
BLUE = RGBColor(31, 95, 174)
GOLD = RGBColor(242, 182, 0)
SKY = RGBColor(235, 243, 255)
LINE = RGBColor(201, 219, 245)
TEXT = RGBColor(38, 55, 78)
MUTED = RGBColor(95, 116, 143)
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
TABLE_INDENT_DXA = CELL_MARGINS_DXA["start"]
MEANINGFUL_COLLECTION_PRIORITY = [
    "books",
    "gifts",
    "home-decor",
    "cookware",
    "pet-assocerries",
    "personal-care",
    "medical-accessories",
    "tools",
    "unique-products",
    "under-35",
]
VALIDATION_COLLECTION_HINTS: dict[str, set[str]] = {
    "Resource Hub": {"books", "gifts", "home-decor", "cookware", "pet-assocerries", "personal-care", "medical-accessories"},
    "Senior Living Guides": {"books", "gifts", "medical-accessories", "gloves"},
    "Home & Living": {"home-decor", "cookware", "tools"},
    "Lifestyle & Wellness": {"books", "personal-care", "medical-accessories", "gloves"},
    "Gift Guides": {"books", "gifts", "home-decor", "cookware"},
    "Home Safety & Organization": {"home-decor", "medical-accessories", "tools", "cookware"},
    "Family & Legacy": {"books", "gifts"},
    "Pet & Home Life": {"pet-assocerries", "gifts", "home-decor"},
}


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def truncate(value: str, limit: int) -> str:
    clean = re.sub(r"\s+", " ", value).strip()
    if len(clean) <= limit:
        return clean
    cut = clean[: max(0, limit - 1)].rsplit(" ", 1)[0].strip()
    return f"{cut}…" if cut else clean[: limit - 1] + "…"


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_product_catalog(path: Path) -> dict[str, Any]:
    payload = load_json(path)
    if isinstance(payload.get("products"), list):
        return payload

    products: list[dict[str, Any]] = []
    for shard in payload.get("shards", []):
        shard_path = path.parent / str(shard.get("file", ""))
        shard_payload = load_json(shard_path)
        products.extend(shard_payload.get("products", []))

    return {**payload, "products": products, "total": payload.get("total", len(products))}


def section_content_width_dxa(section) -> int:
    return (
        int(round(section.page_width.twips))
        - int(round(section.left_margin.twips))
        - int(round(section.right_margin.twips))
    )


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_run_font(run, font_name: str, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    if run._element.rPr is not None:
        r_fonts = run._element.rPr.rFonts
        if r_fonts is not None:
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:cs"), font_name)
            r_fonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    apply_run_font(run, "Arial", size=size if size is not None else 9.5, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int | None = None) -> None:
    widths = [int(width) for width in widths_dxa]
    if not widths:
        raise ValueError("table widths must not be empty")

    total_width = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(total_width))

    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), str(TABLE_INDENT_DXA if indent_dxa is None else indent_dxa))

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for idx, width in enumerate(widths):
        table.columns[idx].width = Twips(width)

    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError("table rows must not contain merged cells")
        row.height = None
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            cell.width = Twips(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(widths[idx]))
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                margin = tc_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), str(CELL_MARGINS_DXA[side]))
                margin.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, font_name, font_size, color in [
        ("Title", "Georgia", 24, NAVY),
        ("Heading 1", "Arial", 18, NAVY),
        ("Heading 2", "Arial", 14, BLUE),
        ("Heading 3", "Arial", 11.5, NAVY),
        ("Heading 4", "Arial", 10.5, NAVY),
    ]:
        style = styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = color

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.font.color.rgb = TEXT


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tblBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tblBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D6E1F2")


def add_paragraph(document: Document, text: str, *, style: str | None = None, color: RGBColor | None = None, bold: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    apply_run_font(run, "Arial", size=10.5, color=color, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12


def add_bullet(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        before, after = text.split(":", 1)
        run = paragraph.add_run(f"{before}:")
        apply_run_font(run, "Arial", size=10.5, bold=True)
        follow = paragraph.add_run(after)
        apply_run_font(follow, "Arial", size=10.5)
    else:
        run = paragraph.add_run(text)
        apply_run_font(run, "Arial", size=10.5)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    apply_run_font(run, "Arial", size=10.5)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12


def add_label_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    header = table.rows[0].cells
    set_cell_text(header[0], "Label", bold=True, color=NAVY, size=10)
    set_cell_text(header[1], "Detail", bold=True, color=NAVY, size=10)
    add_shading(header[0], "EDF4FF")
    add_shading(header[1], "EDF4FF")

    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=BLUE, size=9.5)
        set_cell_text(cells[1], value)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    total_width = section_content_width_dxa(document.sections[0])
    label_width = int(round(total_width * 0.24))
    set_table_geometry(table, [label_width, total_width - label_width])
    document.add_paragraph("")


def add_product_table(document: Document, rows: list[dict[str, str]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    headers = ["Product", "Source signal", "Use case", "Path"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=10)
        add_shading(table.rows[0].cells[idx], "EDF4FF")

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row["product"], bold=True, color=BLUE, size=9.5)
        set_cell_text(cells[1], row["source"], color=MUTED, size=9.25)
        set_cell_text(cells[2], row["use_case"], size=9.25)
        set_cell_text(cells[3], row["path"], color=BLUE, size=9.25)

    total_width = section_content_width_dxa(document.sections[0])
    w0 = int(round(total_width * 0.33))
    w1 = int(round(total_width * 0.17))
    w2 = int(round(total_width * 0.27))
    set_table_geometry(table, [w0, w1, w2, total_width - w0 - w1 - w2])
    document.add_paragraph("")


def add_page_summary_table(document: Document, rows: list[dict[str, str]]) -> None:
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    headers = ["Section", "Page", "Path", "Intent", "Featured products"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=10)
        add_shading(table.rows[0].cells[idx], "EDF4FF")

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row["section"], bold=True, color=BLUE, size=9.25)
        set_cell_text(cells[1], row["title"], size=9.25)
        set_cell_text(cells[2], row["path"], color=BLUE, size=9.0)
        set_cell_text(cells[3], row["intent"], size=9.25)
        set_cell_text(cells[4], row["featured"], size=9.0)

    total_width = section_content_width_dxa(document.sections[0])
    w0 = int(round(total_width * 0.13))
    w1 = int(round(total_width * 0.22))
    w2 = int(round(total_width * 0.22))
    w3 = int(round(total_width * 0.13))
    set_table_geometry(table, [w0, w1, w2, w3, total_width - w0 - w1 - w2 - w3])
    document.add_paragraph("")


def product_reason(product_title: str, page_title: str) -> str:
    title = product_title.lower()
    page = page_title.lower()
    if any(token in title for token in ["planner", "tracker", "journal", "bloom"]):
        return f"Supports the planning and reflection angle behind {page}."
    if any(token in title for token in ["pill", "medication", "medical"]):
        return "Makes medication or essentials easier to keep in one place."
    if any(token in title for token in ["night light", "lamp", "clock"]):
        return "Supports safer evenings and a calmer room at home."
    if any(token in title for token in ["diffuser", "humidifier", "hot water", "hand warmer"]):
        return "Adds a comfort-first self-care layer without feeling overdone."
    if any(token in title for token in ["pet", "dog", "cat", "harness", "brush", "water bottle", "poop bags"]):
        return "Fits the pet-care route and keeps the routine simple."
    if any(token in title for token in ["cookware", "bowl", "cup", "opener", "grinder", "apron"]):
        return "Works well for daily use and giftable household upgrades."
    if any(token in title for token in ["storage", "bag", "stand", "cart", "toilet"]):
        return "Helps the home feel easier to organize and maintain."
    if "walking stick" in title or "cane" in title:
        return "Adds a practical comfort-and-support signal for older adults."
    if any(token in title for token in ["book", "legacy"]):
        return "Reinforces the legacy and memory-preservation theme."
    return f"Strong supporting pick for {page} because it is practical, giftable, and easy to explain."


def build_intro_paragraphs(page: dict[str, Any], kind: str) -> list[str]:
    summary = page["summary"]
    title = page["title"]
    if kind == "hub":
        return [
            "People usually arrive here with a task, not a category name. Pick the card that matches the job you want done, then move to the guide or topic page that feels closest.",
            "The page stays focused on one thing at a time so shoppers can get to the right answer quickly instead of reading through a long, generic directory.",
        ]
    if kind == "category":
        return [
            summary,
            f"This category landing page keeps the browse broad enough to explain the topic, but narrow enough to send readers into the exact subtopic they need next.",
        ]
    return [
        summary,
        f"This topic page is intentionally focused so it can answer the question in plain language, then move the reader toward a matching SALT collection or product.",
    ]


def build_faqs(page: dict[str, Any], kind: str) -> list[tuple[str, str]]:
    title = page["title"]
    if kind == "hub":
        return [
            (
                "What is the Resource Hub for?",
                "It gives shoppers a fast place to start when they have a problem to solve but do not yet know the exact product or guide they need.",
            ),
            (
                "How should someone move through the hub?",
                "Pick the card that matches the task, open the guide, then move into the topic page if the question needs a more specific answer.",
            ),
            (
                "Why does this format work for AEO and GEO?",
                "The pages answer the question first, stay narrow, and connect each answer to real product pages and collections people can click next.",
            ),
        ]
    if kind == "category":
        return [
            (
                f"What does {title} help me solve?",
                f"It turns a broad shopping or planning need into a clearer route, so the reader can move from curiosity to a practical next step.",
            ),
            (
                "Which page should I open after this one?",
                "Open the most relevant subtopic page inside the category, then use the collection link if you want to browse products directly.",
            ),
            (
                "How many products should this page feature?",
                "Three strong examples are enough; that keeps the page useful without turning it into a noisy catalog dump.",
            ),
        ]
    return [
        (
            f"What should a reader do first on {title}?",
            "Start with the smallest useful action: a product, a habit, or a checklist item that makes the problem easier right away.",
        ),
        (
            "How do I keep the page human and not robotic?",
            "Use a calm opening, give one clear answer, and keep the product recommendations tied to an actual everyday need.",
        ),
        (
            "What should this page link to next?",
            "Link to the parent category, one sibling topic, the most relevant collection route, and the first product in the featured set.",
        ),
    ]


def build_ctas(page: dict[str, Any], kind: str) -> list[str]:
    if kind == "hub":
        return [
            "Browse the guide cards first",
            "Open the matching collection next",
            "Contact SALT for support",
        ]
    if kind == "category":
        return [
            "Open the strongest topic page",
            "Browse the matching collection",
            "Return to the Resource Hub",
        ]
    return [
        "Shop the matching collection",
        "Read the next sibling guide",
        "Return to the category page",
    ]


def build_meta_description(page: dict[str, Any], kind: str) -> str:
    base = page["summary"].rstrip(".")
    if kind == "hub":
        extra = " Start with the guide card that matches the task, then move into the topic page if you need a closer answer."
    elif kind == "category":
        extra = " Use the subtopic pages and linked collections to narrow the browse."
    else:
        extra = " Find the related SALT collection and the most useful product picks in one place."
    return truncate(f"{base}.{extra}", 158)


def page_url(category_slug: str | None, topic_slug: str | None = None) -> str:
    if category_slug is None:
        return f"{RESOURCE_BASE}"
    if topic_slug is None:
        return f"{RESOURCE_BASE}/{category_slug}"
    return f"{RESOURCE_BASE}/{category_slug}/{topic_slug}"


def product_url(handle: str) -> str:
    return f"{PRODUCT_BASE}/{handle}"


def flatten_featured_handles(page_specs: list[dict[str, Any]]) -> list[str]:
    handles: list[str] = []
    for page in page_specs:
        for handle in page["featured"]:
            if handle not in handles:
                handles.append(handle)
    return handles


def find_product(products_by_handle: dict[str, dict[str, Any]], handle: str) -> dict[str, Any]:
    if handle not in products_by_handle:
        raise KeyError(f"Product handle not found in Shopify export: {handle}")
    return products_by_handle[handle]


def build_collection_maps(collections_payload: dict[str, Any]) -> tuple[dict[int, list[str]], dict[str, str]]:
    product_collections: dict[int, list[str]] = {}
    collection_titles: dict[str, str] = {}
    for collection_handle, payload in collections_payload.items():
        if not isinstance(payload, dict):
            continue
        collection_titles[collection_handle] = str(payload.get("title", collection_handle))
        for product_id in payload.get("productIds", []):
            product_collections.setdefault(int(product_id), []).append(collection_handle)
    return product_collections, collection_titles


def product_source_label(
    product_id: int,
    best_seller_ids: set[int],
    product_collections: dict[int, list[str]],
    collection_titles: dict[str, str],
) -> str:
    if product_id in best_seller_ids:
        return "Best Sellers collection"

    collection_handles = product_collections.get(product_id, [])
    for preferred_handle in MEANINGFUL_COLLECTION_PRIORITY:
        if preferred_handle in collection_handles:
            return collection_titles.get(preferred_handle, preferred_handle)

    if collection_handles:
        return collection_titles.get(collection_handles[0], collection_handles[0])

    return "Topical collection match"


def validate_featured_products(
    page_specs: list[dict[str, Any]],
    products_by_handle: dict[str, dict[str, Any]],
    best_seller_ids: set[int],
    product_collections: dict[int, list[str]],
) -> None:
    for page in page_specs:
        allowed_collection_handles = VALIDATION_COLLECTION_HINTS.get(page["section"], set())
        if page["kind"] == "hub":
            allowed_collection_handles = VALIDATION_COLLECTION_HINTS["Resource Hub"]
        for handle in page["featured"]:
            product = find_product(products_by_handle, handle)
            if product["id"] in best_seller_ids:
                continue
            available_collections = set(product_collections.get(product["id"], []))
            if not available_collections & allowed_collection_handles:
                raise ValueError(
                    f"Featured product {product['title']} does not match the allowed collection set for {page['section']}"
                )


def build_page_specs() -> list[dict[str, Any]]:
    PLANNER_2 = "the-living-legacy-planner-2nd-edition"
    DAILY_BLOOM = "7-day-mood-mindfulness-tracker"
    MOOD_TRACKER = "7-day-mood-mindfulness-tracker"
    HEALTH_TRACKER = "7-day-health-medication-tracker"
    SOCIAL_TRACKER = "7-day-social-hobby-tracker"
    EXERCISE_PLANNER = "7-day-health-medication-tracker"
    MEAL_PLANNER = "the-living-legacy-planner"
    SLEEP_TRACKER = "7-day-health-medication-tracker"
    ARTHRITIS_GLOVES = "compression-arthritis-gloves-wrist-support-carpal-tunnel-relief"
    PILL_ORG = "7-day-pill-organizer-box-travel-friendly-medicine-dispenser"
    TCARE_PILL = "tcare-travel-pill-organizer-moisture-proof-daily-pill-case-1"
    PILL_BOX_14 = "14-grid-7-day-pill-box-weekly-organizer-for-vitamins-medicine"
    MINI_GPS = "mini-gps-tracker-find-my-app-smart-tag-for-pets-keys"
    LAOPAO_LAMP = "laopao-10w-wireless-charging-led-desk-lamp-dimmable-with-night-light"
    DIGITAL_CLOCK = "digital-wall-clock-time-day-and-temperature-display"
    NONSTICK_9 = "9-piece-nonstick-cookware-set-champagne-lightweight-durable"
    STAINLESS_12 = "12pc-stainless-cookware-set-cook-n-home-kitchen"
    LAPTOP_STAND = "portable-aluminum-laptop-desk-stand-with-mouse-pad"
    MESH_BAGS = "hodr-mesh-bags-lightweight-mesh-stuff-sack-drawstring-storage-bags-compression-pouches-for-camping-hiking-laundry-grocery"
    FOLDING_TOILET = "can-opener-adjustable-stainless-steel-non-slip-manual-jar-bottle-bottle-lid-opener-gadget-home-kitchen-professional-gadgets-tool"
    CAT_HARNESS = "adjustable-cartoon-bee-cat-harness-with-leash-dogs-cats"
    PET_BRUSH = "3-in-1-pet-spray-brush-steam-massage-hair-removal-comb"
    DOG_POOP_BAGS = "dog-poop-bags-10-rolls-portable-pet-waste-bags"
    PET_WATER_BOTTLE = "portable-dog-water-bottle-travel-water-dispenser-for-pets"
    CAT_BED = "cozy-winter-cat-bed-cave-nest-for-cats-small-dogs"
    PET_SEAT_BELT = "adjustable-pet-harness-cat-dog-seat-belt-for-travel"
    DOG_HARNESS = "soft-mesh-dog-harness-breathable-comfort-for-small-dogs-cats"
    PAA_MAA_LAMP = "paamaa-bedside-lamp-night-light-eu-us-plug-led-night-light-ac220v-bedroom-lamp-gift-for-children-cute-night-lamp-for-corridor-wc"
    MINI_NIGHT_LIGHT = "2-pc-led-mini-night-light-switch-plug-in-led-lighting-eye-protection-night-lamp-use-for-bedside-baby-feeding-decoration-bedroom"
    VOLCANO_DIFFUSER = "volcano-flame-aroma-diffuser-360ml-jellyfish-humidifier"
    MULTI_DIFFUSER = "multicolor-humidifier-aromatherapy-flame-diffuser"
    HELLO_KITTY_BOUQUET = "kawaii-hello-kitty-doll-with-artificial-flowers-sanrio-bouquet-gift"
    BRONZE_CANE = "bronze-snake-handle-walking-stick-decorative-cane"
    LEGACY_BOOK = "the-living-legacy-planner"
    MEDICAL_BAG = "portable-weekly-pill-box-8-grids-health-care-travel-organizer"
    VERTICAL_STAND = "vertical-laptop-stand-3-slots-universal-docking-station-for-macbook"
    HOT_WATER_BAG = "hot-water-bottle-bag-warm-belly-hands-feet"
    FOLDING_LUNCH_BOWL = "portable-folding-lunch-box-bowl-sets-silicone-3pcs-set-food-container-outdoor-camping-tableware-set-foldable-salad-bowl-with-lid"
    COLLAPSIBLE_CUP = "outdoors-silicone-folding-cup-with-hanging-hole-creative-water-cup-travel-portable-washing-cup-fashion-travel-silicone-cup"
    ZODIAC_DIFFUSER = "titanic-ship-model-air-humidifier-250ml-essential-oil-diffuser"
    CANDY_DIFFUSER = "aierwill-train-humidifier-ultrasonic-aromatherapy-diffuser"
    CAN_OPENER = "can-opener-adjustable-stainless-steel-non-slip-manual-jar-bottle-bottle-lid-opener-gadget-home-kitchen-professional-gadgets-tool"
    BOTTLE_OPENER = "6-in-1-bottle-opener-multifunctional-screw-cap-jar-can-openers-lid-grip-opener-home-camping-safety-can-opener-kitchen-gadgets"
    FOOD_SPRAYER = "atwfs-electric-wireless-bottle-pump"

    return [
        {
            "section": "Resource Hub",
            "kind": "hub",
            "title": "Resource Hub",
            "slug": "resources",
            "summary": "Pick the question you're trying to answer, then jump into the guide or topic page that fits it best.",
            "collection_route": None,
            "featured": [PLANNER_2, DAILY_BLOOM, DIGITAL_CLOCK],
        },
        {
            "section": "Senior Living Guides",
            "kind": "category",
            "title": "Senior Living Guides",
            "slug": "senior-living-guides",
            "summary": "Practical guides for adults who want safer routines, gentler support, and thoughtful products that make daily life easier.",
            "collection_route": "/collections/senior-living-solutions",
            "featured": [PLANNER_2, PILL_ORG, ARTHRITIS_GLOVES],
            "topics": [
                {
                    "title": "Best Gifts for Seniors",
                    "summary": "A thoughtful guide to gifts that feel useful, comfortable, and respectful, with products that support everyday routines instead of collecting dust.",
                    "collection_route": "/collections/gifts",
                    "featured": [PLANNER_2, PILL_ORG, BRONZE_CANE],
                },
                {
                    "title": "Home Safety Tips",
                    "summary": "A calm starting point for safer rooms, gentler night-time navigation, and small upgrades that reduce avoidable friction at home.",
                    "collection_route": "/collections/home-decor-lighting",
                    "featured": [MINI_NIGHT_LIGHT, DIGITAL_CLOCK, MEDICAL_BAG],
                },
                {
                    "title": "Caregiver Resources",
                    "summary": "A practical guide for caregivers who need products and systems that save time, reduce stress, and keep important details organized.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [HEALTH_TRACKER, PILL_BOX_14, MEDICAL_BAG],
                },
            ],
        },
        {
            "section": "Home & Living",
            "kind": "category",
            "title": "Home & Living",
            "slug": "home-living",
            "summary": "Calm, room-by-room ideas for organizing small spaces, reducing clutter, and making the home feel more settled.",
            "collection_route": "/collections/home-kitchen",
            "featured": [LAOPAO_LAMP, DIGITAL_CLOCK, MESH_BAGS],
            "topics": [
                {
                    "title": "How to Stay Organized at Home",
                    "summary": "A simple guide for turning clutter into a manageable system, one room and one habit at a time.",
                    "collection_route": "/collections/home-kitchen",
                    "featured": [MESH_BAGS, PLANNER_2, DIGITAL_CLOCK],
                },
                {
                    "title": "Small Space Organization Tips",
                    "summary": "Space-saving ideas for apartments, shared rooms, and compact homes where every inch has to work a little harder.",
                    "collection_route": "/collections/home-kitchen",
                    "featured": [FOLDING_LUNCH_BOWL, COLLAPSIBLE_CUP, LAPTOP_STAND],
                },
                {
                    "title": "Decluttering Your Home",
                    "summary": "A straightforward plan for clearing visual noise, keeping what matters, and making it easier to maintain order afterward.",
                    "collection_route": "/collections/home-kitchen",
                    "featured": [FOLDING_TOILET, MESH_BAGS, DIGITAL_CLOCK],
                },
                {
                    "title": "Creating a Comfortable Living Space",
                    "summary": "Ideas for adding warmth, light, and small comforts so the home feels inviting without becoming busy.",
                    "collection_route": "/collections/home-decor-lighting",
                    "featured": [LAOPAO_LAMP, PAA_MAA_LAMP, VOLCANO_DIFFUSER],
                },
            ],
        },
        {
            "section": "Lifestyle & Wellness",
            "kind": "category",
            "title": "Lifestyle & Wellness",
            "slug": "lifestyle-wellness",
            "summary": "Simple, sustainable habits and comfort-forward products that help the day feel less rushed and more manageable.",
            "collection_route": "/collections/health-wellness",
            "featured": [DAILY_BLOOM, MOOD_TRACKER, SLEEP_TRACKER],
            "topics": [
                {
                    "title": "Simple Habits for a Less Stressful Life",
                    "summary": "A gentle reset for people who want less noise in the day and more rhythm without turning life into a project.",
                    "collection_route": "/collections/health-wellness",
                    "featured": [DAILY_BLOOM, MOOD_TRACKER, SLEEP_TRACKER],
                },
                {
                    "title": "Creating Better Daily Routines",
                    "summary": "A page for building repeatable routines that make mornings, evenings, and in-between moments feel easier to trust.",
                    "collection_route": "/collections/health-wellness",
                    "featured": [PLANNER_2, SOCIAL_TRACKER, EXERCISE_PLANNER],
                },
                {
                    "title": "Work-Life Balance Tips",
                    "summary": "Practical ways to protect focus and personal time so work does not crowd out the rest of the day.",
                    "collection_route": "/collections/health-wellness",
                    "featured": [LAPTOP_STAND, DIGITAL_CLOCK, DAILY_BLOOM],
                },
                {
                    "title": "Self-Care at Home",
                    "summary": "Low-effort comfort ideas for the moments when you need rest, a pause, or a more soothing environment.",
                    "collection_route": "/collections/health-wellness",
                    "featured": [MULTI_DIFFUSER, HOT_WATER_BAG, SLEEP_TRACKER],
                },
            ],
        },
        {
            "section": "Gift Guides",
            "kind": "category",
            "title": "Gift Guides",
            "slug": "gift-guides",
            "summary": "Occasion-led gift ideas that feel personal, useful, and easy to choose when you want the present to be appreciated, not overthought.",
            "collection_route": "/collections/gifts",
            "featured": [DAILY_BLOOM, PLANNER_2, HELLO_KITTY_BOUQUET],
            "topics": [
                {
                    "title": "Best Gifts for Mom",
                    "summary": "Warm, practical gifts for moms who appreciate something useful, thoughtful, and a little more personal than the usual default.",
                    "collection_route": "/collections/gifts",
                    "featured": [DAILY_BLOOM, LAOPAO_LAMP, NONSTICK_9],
                },
                {
                    "title": "Best Gifts for Dad",
                    "summary": "Useful, giftable picks for dads who value function, good design, and items they can actually use.",
                    "collection_route": "/collections/gifts",
                    "featured": [VERTICAL_STAND, STAINLESS_12, CAN_OPENER],
                },
                {
                    "title": "Best Gifts for Grandparents",
                    "summary": "Comfort-first gift ideas for grandparents that feel kind, easy to enjoy, and simple to incorporate into daily life.",
                    "collection_route": "/collections/gifts",
                    "featured": [PLANNER_2, PILL_ORG, BRONZE_CANE],
                },
                {
                    "title": "Housewarming Gift Ideas",
                    "summary": "A curated mix of gifts that help a new home feel finished, comfortable, and ready to use.",
                    "collection_route": "/collections/home-decor-lighting",
                    "featured": [NONSTICK_9, LAOPAO_LAMP, DIGITAL_CLOCK],
                },
                {
                    "title": "Holiday Gift Guides",
                    "summary": "Seasonal gift ideas that keep shopping organized and help you choose something meaningful without the last-minute scramble.",
                    "collection_route": "/collections/gifts",
                    "featured": [DAILY_BLOOM, HELLO_KITTY_BOUQUET, MULTI_DIFFUSER],
                },
            ],
        },
        {
            "section": "Home Safety & Organization",
            "kind": "category",
            "title": "Home Safety & Organization",
            "slug": "home-safety-organization",
            "summary": "Straightforward guidance for keeping the home safer, important items easier to find, and routines easier to maintain.",
            "collection_route": "/collections/home-kitchen",
            "featured": [MINI_NIGHT_LIGHT, DIGITAL_CLOCK, MEDICAL_BAG],
            "topics": [
                {
                    "title": "Home Safety Tips for Every Age",
                    "summary": "A family-friendly guide to safer home details that matter at every stage of life, from lighting to organization.",
                    "collection_route": "/collections/home-decor-lighting",
                    "featured": [MINI_NIGHT_LIGHT, DIGITAL_CLOCK, MINI_GPS],
                },
                {
                    "title": "Organizing Important Documents",
                    "summary": "A guide for keeping critical papers, passwords, and emergency information together so they are easier to find when needed.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [PLANNER_2, DAILY_BLOOM, HEALTH_TRACKER],
                },
                {
                    "title": "Family Emergency Preparedness",
                    "summary": "A practical checklist for making sure the household has the basics ready before a problem turns urgent.",
                    "collection_route": "/collections/travel-outdoor",
                    "featured": [MEDICAL_BAG, MINI_GPS, FOLDING_TOILET],
                },
                {
                    "title": "Keeping Your Home Clutter-Free",
                    "summary": "A maintenance-minded page for keeping surfaces clear, storage simple, and the house easier to reset each day.",
                    "collection_route": "/collections/home-kitchen",
                    "featured": [MESH_BAGS, FOLDING_LUNCH_BOWL, COLLAPSIBLE_CUP],
                },
            ],
        },
        {
            "section": "Family & Legacy",
            "kind": "category",
            "title": "Family & Legacy",
            "slug": "family-legacy",
            "summary": "Thoughtful planning pages that help families preserve memories, organize key information, and make the future feel less scattered.",
            "collection_route": "/collections/senior-living-solutions",
            "featured": [PLANNER_2, LEGACY_BOOK, DAILY_BLOOM],
            "topics": [
                {
                    "title": "Preserving Family Memories",
                    "summary": "A gentle guide to saving stories, notes, and keepsakes so family history stays available and meaningful.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [LEGACY_BOOK, DAILY_BLOOM, PLANNER_2],
                },
                {
                    "title": "Why Every Family Should Have Important Information Organized",
                    "summary": "A clear case for keeping essential household and care information in one place before life gets complicated.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [PLANNER_2, HEALTH_TRACKER, PILL_BOX_14],
                },
                {
                    "title": "Creating a Family Legacy",
                    "summary": "A page for capturing values, stories, and practical details that outlast a single season.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [LEGACY_BOOK, DAILY_BLOOM, SOCIAL_TRACKER],
                },
                {
                    "title": "Planning for the Future",
                    "summary": "A calm, practical page that helps families think ahead without making the conversation feel heavy.",
                    "collection_route": "/collections/senior-living-solutions",
                    "featured": [PLANNER_2, EXERCISE_PLANNER, MEAL_PLANNER],
                },
            ],
        },
        {
            "section": "Pet & Home Life",
            "kind": "category",
            "title": "Pet & Home Life",
            "slug": "pet-home-life",
            "summary": "Practical pet-care pages that keep feeding, travel, grooming, and everyday home routines simpler for people and pets.",
            "collection_route": "/collections/pet-essentials",
            "featured": [CAT_HARNESS, PET_BRUSH, PET_WATER_BOTTLE],
            "topics": [
                {
                    "title": "Organizing Pet Supplies",
                    "summary": "A tidy, practical guide for keeping leashes, grooming tools, waste bags, and feeding gear easy to reach.",
                    "collection_route": "/collections/pet-essentials",
                    "featured": [DOG_POOP_BAGS, PET_BRUSH, PET_SEAT_BELT],
                },
                {
                    "title": "Making Your Home Pet-Friendly",
                    "summary": "Simple ideas for sharing the home with pets in a way that feels comfortable, safe, and easy to maintain.",
                    "collection_route": "/collections/pet-essentials",
                    "featured": [CAT_BED, DOG_HARNESS, CAT_HARNESS],
                },
                {
                    "title": "Travel Tips for Pet Owners",
                    "summary": "A handy planning guide for road trips, appointments, and overnight stays with pets, plus the gear that makes travel smoother.",
                    "collection_route": "/collections/travel-outdoor",
                    "featured": [PET_WATER_BOTTLE, MINI_GPS, PET_SEAT_BELT],
                },
            ],
        },
    ]


def build_document(
    page_specs: list[dict[str, Any]],
    products_by_handle: dict[str, dict[str, Any]],
    best_seller_ids: set[int],
    product_collections: dict[int, list[str]],
    collection_titles: dict[str, str],
) -> Document:
    document = Document()
    style_document(document)

    # Cover
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SALT AEO/GEO Resource Hub Content Package")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Word source of truth for the Resource Hub, category pages, topic pages, product shortlist, and homepage-style guidance.")
    apply_run_font(run, "Arial", size=11.5, color=MUTED)

    stamp = document.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    apply_run_font(run, "Arial", size=9.5, color=MUTED)

    document.add_paragraph("")
    add_paragraph(document, "Source files used", color=BLUE, bold=True)
    add_bullet(document, "public/data/products.json")
    add_bullet(document, "public/data/collection-products.json")
    add_bullet(document, "src/lib/site-navigation.ts")
    add_bullet(document, "src/pages/HomePage.tsx")
    add_bullet(document, "src/index.css and index.html")

    document.add_page_break()

    # Strategy section
    document.add_heading("Strategy Overview", level=1)
    add_paragraph(
        document,
        "This package is built around answer-first pages that help AI search engines understand the page intent, while still feeling like SALT: calm, practical, and human.",
    )
    add_paragraph(
        document,
        "The content uses the existing Shopify export as the product source of truth. Best Sellers act as the ranking proxy, and topical collection matches fill any gaps where the best seller list is not enough.",
    )
    add_label_value_table(
        document,
        [
            ("Content model", "1 hub page, 7 category pages, 27 supporting topic pages"),
            ("Primary product source", "public/data/products.json"),
            ("Ranking proxy", "public/data/collection-products.json Best Sellers collection"),
            ("Page intent", "AEO / GEO visibility with direct shopping support"),
            ("Brand voice", "Calm, warm, practical, and human"),
        ],
    )

    document.add_heading("Homepage Style Audit", level=1)
    add_paragraph(
        document,
        "The SALT homepage already signals a strong visual and editorial direction: light mode only, soft blue and gold accents, airy cards, rounded geometry, and a direct path from inspiration to product.",
    )
    audit_rows = [
        ("Color system", "Keep the light-blue base, navy text, and gold accent touches. Avoid dark-mode fallback or high-contrast neon styling."),
        ("Typography", "Use clean, readable sans-serif body copy with a restrained serif or display accent for section headings."),
        ("Layout rhythm", "Keep pages spacious, with one strong intro, one product block, then short support sections instead of dense walls of copy."),
        ("CTA behavior", "Use one primary action and two secondary actions. Make the next step obvious, but do not overwhelm the reader."),
        ("Card treatment", "Stick with rounded cards, soft borders, and subtle shadow depth. Avoid sharp edges and heavy chrome."),
        ("Editorial tone", "Write as if a helpful merchandiser is guiding the shopper, not as if the page is trying to game search engines."),
    ]
    add_label_value_table(document, audit_rows)

    document.add_heading("AEO / GEO Page Rules", level=1)
    for rule in [
        "Lead with the answer in the first paragraph and keep the main search intent visible in the H1.",
        "Use one clear page theme per URL. If a page has two jobs, split it into two pages.",
        "Keep the FAQ section short: three questions is enough when the answers are direct and useful.",
        "Add three real products to each page, chosen from the Shopify export rather than generic placeholders.",
        "Link each page to its parent category, one sibling guide, and one relevant collection or product page.",
        "Keep the metadata human. The title should read well in search results, not look like a keyword pile.",
        "Use natural language variants of the main term so AI search can map the page to a broader query family.",
        "Avoid over-explaining. The reader should get the answer quickly and then know exactly what to do next.",
    ]:
        add_bullet(document, rule)

    document.add_heading("Product Shortlist", level=1)
    add_paragraph(
        document,
        "These products are the strongest editorial matches for the hub package. The shortlist prefers Best Sellers when they fit the topic, then falls back to topical collection matches.",
    )

    unique_handles = flatten_featured_handles(page_specs)
    unique_products: list[dict[str, Any]] = [find_product(products_by_handle, handle) for handle in unique_handles]
    unique_products.sort(
        key=lambda product: (
            0 if product["id"] in best_seller_ids else 1,
            product["title"].lower(),
        )
    )

    product_rows: list[dict[str, str]] = []
    for product in unique_products:
        source = product_source_label(product["id"], best_seller_ids, product_collections, collection_titles)
        page_titles = [page["title"] for page in page_specs if product["handle"] in page["featured"]]
        use_case = product_reason(product["title"], page_titles[0] if page_titles else product["title"])
        product_rows.append(
            {
                "product": product["title"],
                "source": source,
                "use_case": use_case,
                "path": product_url(product["handle"]),
            }
        )

    add_product_table(document, product_rows[:24])

    document.add_heading("Page Matrix", level=1)
    add_paragraph(
        document,
        "The matrix below shows the final page hierarchy. Paths are written for the resource hub route structure that SALT can publish directly.",
    )

    summary_rows: list[dict[str, str]] = []
    for category in page_specs:
        kind = category["kind"]
        if kind == "hub":
            summary_rows.append(
                {
                    "section": category["section"],
                    "title": category["title"],
                    "path": page_url(None),
                    "intent": "Root answer-first landing page",
                    "featured": ", ".join(find_product(products_by_handle, handle)["title"] for handle in category["featured"]),
                }
            )
            continue

        summary_rows.append(
            {
                "section": category["section"],
                "title": category["title"],
                "path": page_url(category["slug"]),
                "intent": "Category landing page",
                "featured": ", ".join(find_product(products_by_handle, handle)["title"] for handle in category["featured"]),
            }
        )
        for topic in category.get("topics", []):
            summary_rows.append(
                {
                    "section": category["section"],
                    "title": topic["title"],
                    "path": page_url(category["slug"], topic["slug"]),
                    "intent": "Supporting topic page",
                    "featured": ", ".join(find_product(products_by_handle, handle)["title"] for handle in topic["featured"]),
                }
            )

    add_page_summary_table(document, summary_rows)

    document.add_page_break()

    # Detailed pages
    for page in page_specs:
        kind = page["kind"]
        if kind == "hub":
            document.add_heading(page["title"], level=1)
            add_paragraph(document, page["summary"])
            add_label_value_table(
                document,
                [
                    ("Type", "Hub page"),
                    ("Path", page_url(None)),
                    ("SEO title", f"{page['title']} | SALT Resource Hub"),
                    ("Meta description", build_meta_description(page, kind)),
                    ("Primary intent", "Answer-first entry point for the full content cluster"),
                ],
            )
            document.add_heading("Choose Your Path", level=2)
            add_paragraph(
                document,
                "The hub works best when it starts from the problem the shopper has. Each category card below points to a guide first, then to the more specific topic pages.",
            )
            for guide in page_specs[1:]:
                add_bullet(
                    document,
                    f"{guide['title']}: {guide['summary']} Open the guide first, then jump into the most relevant topic page.",
                )
        else:
            document.add_heading(page["section"], level=1)
            add_paragraph(document, page["summary"], color=TEXT)
            add_label_value_table(
                document,
                [
                    ("Type", "Category landing page" if kind == "category" else "Topic page"),
                    ("Path", page_url(page.get("slug")) if kind == "category" else page_url(page["category_slug"], page["slug"])),
                    ("SEO title", f"{page['title']} | SALT Resource Hub"),
                    ("Meta description", build_meta_description(page, kind)),
                    ("Primary intent", page.get("primary_intent", "Answer-first SALT guide")),
                ],
            )

        if kind == "hub":
            intro_paragraphs = build_intro_paragraphs(page, kind)
            for paragraph in intro_paragraphs:
                add_paragraph(document, paragraph)

            document.add_heading("Choose your path", level=2)
            add_paragraph(
                document,
                "The hub works best when it starts from the problem the shopper has. Each category card points to a guide first, then to the more specific topic pages.",
            )

            document.add_heading("FAQ Prompts", level=2)
            for question, answer in build_faqs(page, kind):
                add_bullet(document, f"Q: {question} A: {answer}")

            document.add_heading("Call to Action", level=2)
            for cta in build_ctas(page, kind):
                add_bullet(document, cta)

            document.add_heading("Internal Links", level=2)
            for link in [
                f"Senior Living Guides: {page_url('senior-living-guides')}",
                f"Home & Living: {page_url('home-living')}",
                f"Lifestyle & Wellness: {page_url('lifestyle-wellness')}",
                f"Gift Guides: {page_url('gift-guides')}",
                f"Home Safety & Organization: {page_url('home-safety-organization')}",
                f"Family & Legacy: {page_url('family-legacy')}",
                f"Pet & Home Life: {page_url('pet-home-life')}",
            ]:
                add_bullet(document, link)

            document.add_heading("Featured Products", level=2)
            for handle in page["featured"]:
                product = find_product(products_by_handle, handle)
                add_bullet(
                    document,
                    f"{product['title']} - {product_reason(product['title'], page['title'])} {product_url(product['handle'])}",
                )

            document.add_page_break()
            continue

        document.add_heading("Intro Copy", level=2)
        for paragraph in build_intro_paragraphs(page, kind):
            add_paragraph(document, paragraph)

        document.add_heading("FAQ Prompts", level=2)
        for question, answer in build_faqs(page, kind):
            add_bullet(document, f"Q: {question} A: {answer}")

        document.add_heading("Call to Action", level=2)
        for cta in build_ctas(page, kind):
            add_bullet(document, cta)

        document.add_heading("Internal Links", level=2)
        category_slug = page.get("category_slug")
        if kind == "category":
            for link in [
                f"Resource Hub: {page_url(None)}",
                f"{page['title']} collection route: {page['collection_route']}",
            ]:
                add_bullet(document, link)
            for topic in page["topics"]:
                add_bullet(document, f"{topic['title']}: {page_url(page['slug'], topic['slug'])}")
        else:
            sibling_topic = None
            for sibling in next(cat for cat in page_specs if cat["slug"] == category_slug)["topics"]:
                if sibling["slug"] != page["slug"]:
                    sibling_topic = sibling
                    break
            link_rows = [
                f"Parent category: {page_url(category_slug)}",
                f"Matching collection: {page['collection_route']}",
            ]
            if sibling_topic:
                link_rows.append(f"Sibling guide: {page_url(category_slug, sibling_topic['slug'])}")
            first_product = find_product(products_by_handle, page["featured"][0])
            link_rows.append(f"Featured product: {product_url(first_product['handle'])}")
            for link in link_rows:
                add_bullet(document, link)

        document.add_heading("Featured Products", level=2)
        for handle in page["featured"]:
            product = find_product(products_by_handle, handle)
            add_bullet(
                document,
                f"{product['title']} - {product_reason(product['title'], page['title'])} {product_url(product['handle'])}",
            )

        if kind == "category":
            document.add_heading("Subtopic Map", level=2)
            for topic in page["topics"]:
                add_numbered(
                    document,
                    f"{topic['title']} - {page_url(page['slug'], topic['slug'])}",
                )
        document.add_page_break()

    # Publishing guidance
    document.add_heading("Publishing Checklist", level=1)
    for item in [
        "Use the exact title and path for each page so the hub hierarchy stays obvious to search engines.",
        "Keep the first paragraph answer-first and make sure the keyword appears naturally, not repeatedly.",
        "Add FAQ schema when the page goes live so search engines can extract the question-answer pairs cleanly.",
        "Link back to the hub and the category page from each topic page to preserve the topical cluster.",
        "Place the three featured products near the intro and keep the shopping path close to the content.",
        "Review the page in light mode and ensure the visual hierarchy still feels like SALT: airy, warm, and practical.",
    ]:
        add_bullet(document, item)

    return document


def main() -> int:
    products_payload = load_product_catalog(PRODUCTS_PATH)
    collections_payload = load_json(COLLECTION_PRODUCTS_PATH)

    products = products_payload["products"]
    products_by_handle = {product["handle"]: product for product in products}
    collections = collections_payload["collections"]
    best_seller_ids = set(collections["appplaza-best-sellers"]["productIds"])
    product_collections, collection_titles = build_collection_maps(collections)

    categories = build_page_specs()
    page_specs: list[dict[str, Any]] = [
        {
            "kind": categories[0]["kind"],
            "section": categories[0]["section"],
            "title": categories[0]["title"],
            "slug": categories[0]["slug"],
            "summary": categories[0]["summary"],
            "featured": categories[0]["featured"],
        }
    ]
    for category in categories[1:]:
        page_specs.append(
            {
                "kind": category["kind"],
                "section": category["section"],
                "title": category["title"],
                "slug": category["slug"],
                "summary": category["summary"],
                "collection_route": category["collection_route"],
                "featured": category["featured"],
                "topics": [
                    {
                        "title": topic["title"],
                        "slug": slugify(topic["title"]),
                        "summary": topic["summary"],
                        "collection_route": topic["collection_route"],
                        "featured": topic["featured"],
                    }
                    for topic in category.get("topics", [])
                ],
            }
        )

    validate_featured_products(page_specs, products_by_handle, best_seller_ids, product_collections)

    document = build_document(page_specs, products_by_handle, best_seller_ids, product_collections, collection_titles)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
